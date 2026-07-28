from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


OUTPUT = Path(
    r"C:\Users\dimah\IdeaProjects\Tarification\docs"
    r"\Инструкция_по_модулю_Педагогические_советы.docx"
)
SCREENSHOT_DIR = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "screenshots"
    / "pedagogical-councils"
)

# compact_reference_guide + named "Russian office A4" override.
PAGE_WIDTH_DXA = 11906
PAGE_HEIGHT_DXA = 16838
MARGIN_DXA = 1134
CONTENT_WIDTH_DXA = 9638
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120

FONT = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "111111"
MUTED = "626B75"
CAUTION = "7A5A00"
RISK = "9B1C1C"
POSITIVE = "1F3A5F"
BORDER = "B8C4D1"


def set_run_font(
    run,
    *,
    name: str = FONT,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_keep(paragraph, *, with_next: bool = False, together: bool = False) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if with_next:
        ppr.append(OxmlElement("w:keepNext"))
    if together:
        ppr.append(OxmlElement("w:keepLines"))
    ppr.append(OxmlElement("w:widowControl"))


def shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell,
    *,
    top: int = CELL_TOP_BOTTOM_DXA,
    bottom: int = CELL_TOP_BOTTOM_DXA,
    start: int = CELL_START_END_DXA,
    end: int = CELL_START_END_DXA,
) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Mm(widths[index] * 25.4 / 1440)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[index]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def prevent_row_split(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def configure_cell_paragraph(paragraph, *, size: float = 10.2, color: str = BLACK) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_keep(paragraph, together=True)
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10.2,
    color: str = BLACK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    set_paragraph_keep(paragraph, together=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_text_paragraph(
    container,
    text: str = "",
    *,
    bold_lead: str | None = None,
    size: float = 11,
    color: str = BLACK,
    italic: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.25,
    keep_together: bool = False,
) -> object:
    paragraph = container.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    set_paragraph_keep(paragraph, together=keep_together)
    if bold_lead and text.startswith(bold_lead):
        lead_run = paragraph.add_run(bold_lead)
        set_run_font(lead_run, size=size, bold=True, color=color)
        rest_run = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest_run, size=size, italic=italic, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, italic=italic, color=color)
    return paragraph


def add_heading(doc, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    set_paragraph_keep(paragraph, with_next=True, together=True)
    return paragraph


def add_page_heading(doc, section_number: str, title: str, subtitle: str | None = None) -> None:
    kicker = add_text_paragraph(
        doc,
        section_number.upper(),
        size=9.5,
        color=BLUE,
        before=0,
        after=2,
        line_spacing=1.0,
        keep_together=True,
    )
    kicker.runs[0].bold = True
    kicker.runs[0].font.all_caps = True
    heading = add_heading(doc, title, level=1)
    if subtitle:
        add_text_paragraph(
            doc,
            subtitle,
            size=10.5,
            color=MUTED,
            after=10,
            line_spacing=1.15,
            keep_together=True,
        )


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])
    set_run_font(run, size=9, color=MUTED)


def configure_page(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0


def add_numbering_abstract(document: Document, *, decimal: bool) -> int:
    numbering = document.part.numbering_part.element
    existing_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    abstract_id = max(existing_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if decimal else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if decimal else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    ppr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)

    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)
    rpr.append(rfonts)
    level.append(rpr)
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def create_num_instance(document: Document, abstract_id: int) -> int:
    numbering = document.part.numbering_part.element
    existing_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(
    document: Document,
    text: str,
    *,
    num_id: int,
    bold_lead: str | None = None,
    color: str = BLACK,
) -> object:
    paragraph = document.add_paragraph()
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_paragraph_keep(paragraph, together=True)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=color)
    return paragraph


def add_bullets(document: Document, items: Iterable[str], bullet_abstract_id: int) -> None:
    num_id = create_num_instance(document, bullet_abstract_id)
    for item in items:
        add_list_item(document, item, num_id=num_id)


def add_numbered_steps(
    document: Document,
    items: Iterable[tuple[str, str]],
    decimal_abstract_id: int,
) -> None:
    num_id = create_num_instance(document, decimal_abstract_id)
    for title, detail in items:
        text = f"{title}. {detail}"
        add_list_item(document, text, num_id=num_id, bold_lead=f"{title}.")


def add_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    *,
    font_size: float = 10.2,
    first_col_bold: bool = False,
) -> object:
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, text in enumerate(headers):
        set_cell_shading(header.cells[index], LIGHT_BLUE)
        set_cell_text(header.cells[index], text, bold=True, size=10, color=INK_BLUE)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_values):
            set_cell_text(
                row.cells[index],
                value,
                bold=first_col_bold and index == 0,
                size=font_size,
            )
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(3)
    return table


def add_callout(
    document: Document,
    label: str,
    text: str,
    *,
    fill: str = CALLOUT,
    label_color: str = DARK_BLUE,
) -> object:
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=label_color, size=8)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_keep(paragraph, together=True)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, bold=True, color=label_color)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=BLACK)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return table


def add_code_sample(document: Document, lines: Sequence[str], title: str) -> None:
    caption = add_text_paragraph(
        document,
        title,
        size=10,
        color=DARK_BLUE,
        before=2,
        after=4,
        keep_together=True,
    )
    caption.runs[0].bold = True
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="CBD5E1", size=6)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.12
    set_paragraph_keep(paragraph, together=True)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9.2, color=BLACK)
        if index < len(lines) - 1:
            run.add_break()
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_figure(
    document: Document,
    image_name: str,
    caption_text: str,
    alt_text: str,
    *,
    width_mm: float = 168,
) -> None:
    image_path = SCREENSHOT_DIR / image_name
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    set_paragraph_keep(paragraph, with_next=True, together=True)

    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))
    for doc_pr in run._r.iter(qn("wp:docPr")):
        doc_pr.set("descr", alt_text)
        doc_pr.set("title", caption_text)

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.05
    set_paragraph_keep(caption, together=True)
    caption_run = caption.add_run(caption_text)
    set_run_font(caption_run, size=9.2, italic=True, color=MUTED)


def add_footer_header(section) -> None:
    section.different_first_page_header_footer = True

    first_header = section.first_page_header
    first_header.paragraphs[0].clear()
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].clear()
    fp = first_footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Инструкция пользователя • Версия 1.0 • 28.07.2026")
    set_run_font(r, size=8.5, color=MUTED)

    header = section.header
    header.paragraphs[0].clear()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("Модуль «Педагогические советы» • Инструкция пользователя")
    set_run_font(r, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    footer.paragraphs[0].clear()
    fp2 = footer.paragraphs[0]
    fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp2.paragraph_format.space_before = Pt(0)
    r = fp2.add_run("Страница ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp2, "PAGE")
    r = fp2.add_run(" из ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp2, "NUMPAGES")


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    configure_page(section)
    configure_styles(document)
    add_footer_header(section)

    props = document.core_properties
    props.title = "Инструкция по работе с модулем «Педагогические советы»"
    props.subject = "Создание протоколов, приложений и выписок"
    props.author = "Система управления школой"
    props.keywords = "педагогический совет, протокол, выписка, инструкция"
    props.comments = "Актуально на 28.07.2026"

    decimal_abstract = add_numbering_abstract(document, decimal=True)
    bullet_abstract = add_numbering_abstract(document, decimal=False)

    # Cover page: editorial_cover pattern, compact_reference_guide preset.
    cover_label = add_text_paragraph(
        document,
        "ИНСТРУКЦИЯ ПОЛЬЗОВАТЕЛЯ",
        size=10,
        color=BLUE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=92,
        after=18,
        line_spacing=1.0,
        keep_together=True,
    )
    cover_label.runs[0].bold = True
    cover_label.runs[0].font.all_caps = True

    title = add_text_paragraph(
        document,
        "Модуль\n«Педагогические советы»",
        size=28,
        color=INK_BLUE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
        line_spacing=1.05,
        keep_together=True,
    )
    for run in title.runs:
        run.bold = True

    add_text_paragraph(
        document,
        "Создание протоколов, работа с приложениями,\n"
        "выпуск документов и подготовка выписок",
        size=14,
        color=DARK_BLUE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=36,
        line_spacing=1.18,
        keep_together=True,
    )
    add_text_paragraph(
        document,
        "Для директора, заместителей директора, методистов и секретарей педагогического совета",
        size=10.5,
        color=MUTED,
        italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=80,
        line_spacing=1.2,
        keep_together=True,
    )
    add_text_paragraph(
        document,
        "Версия 1.0",
        size=11,
        color=INK_BLUE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=2,
        keep_together=True,
    ).runs[0].bold = True
    add_text_paragraph(
        document,
        "Актуально на 28 июля 2026 года",
        size=10,
        color=MUTED,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
        keep_together=True,
    )

    add_page_break(document)

    # 1. Quick start.
    add_page_heading(
        document,
        "Раздел 1",
        "Быстрый старт",
        "Весь рабочий цикл — от формы до готового Word-документа.",
    )
    add_callout(
        document,
        "Главный принцип",
        "Секретарь один раз заполняет форму протокола. На основании этих данных система "
        "формирует полный протокол и любые выписки из него — повторно набирать текст не нужно.",
        fill=LIGHT_BLUE,
    )
    add_heading(document, "Рабочий маршрут", level=2)
    add_numbered_steps(
        document,
        [
            ("Выберите учебный год", "Откройте «Документы» → «Педагогические советы» и укажите год."),
            ("Создайте протокол", "Заполните реквизиты заседания, председателя и секретаря."),
            ("Добавьте пункты", "Для каждого вопроса укажите докладчика, содержание, решение и голоса."),
            ("Прикрепите приложения", "Загрузите готовые файлы Word к нужным пунктам."),
            ("Сохраните и выпустите", "Проверьте предварительный просмотр и нажмите нужную кнопку."),
            ("Скачайте документ", "Получите полный протокол или выписку по выбранным пунктам."),
        ],
        decimal_abstract,
    )
    add_heading(document, "Что делает система автоматически", level=2)
    add_bullets(
        document,
        [
            "подставляет наименование и бланк школы согласно настройке сервера;",
            "берёт сотрудников и основные должности из раздела кадров;",
            "показывает свободный номер приложения и связывает файл с пунктом;",
            "собирает протокол и выбранные приложения в один Word-файл;",
            "добавляет в выписку приложения именно к выбранным пунктам;",
            "ставит заверителя выписки из данных пользователя, который её формирует.",
        ],
        bullet_abstract,
    )
    add_callout(
        document,
        "Подписание",
        "Система готовит Word-файлы для печати. Подписи председателя, секретаря, заверителей "
        "и утверждающего ставятся на бумаге.",
        fill=CALLOUT,
    )

    add_page_break(document)

    # 2. Access and protocol list.
    add_page_heading(
        document,
        "Раздел 2",
        "Вход в модуль и список протоколов",
        "Доступные кнопки зависят от прав, назначенных пользователю.",
    )
    add_heading(document, "Как открыть модуль", level=2)
    add_numbered_steps(
        document,
        [
            ("Откройте главное меню", "Выберите раздел «Документы»."),
            ("Перейдите к советам", "Нажмите «Педагогические советы»."),
            ("Выберите год", "В поле «Учебный год» укажите нужный период, например 2025/2026."),
        ],
        decimal_abstract,
    )
    add_callout(
        document,
        "Учебный год",
        "В системе учебный год длится с 1 августа первого года по 31 июля следующего. "
        "Например, для 2025/2026 допустимы даты с 01.08.2025 по 31.07.2026.",
        fill=CALLOUT,
    )
    add_heading(document, "Права пользователя", level=2)
    add_table(
        document,
        ["Право", "Что доступно"],
        [
            ("Просмотр", "Открыть модуль и увидеть список протоколов."),
            ("Редактирование", "Создать, открыть, изменить, выпустить, перевыпустить и удалить протокол."),
            ("Импорт", "Загрузить старый готовый протокол в формате Word .docx."),
            ("Экспорт", "Скачать полный протокол, приложения и сформировать выписку."),
        ],
        [2700, 6938],
        first_col_bold=True,
    )
    add_text_paragraph(
        document,
        "Если нужной кнопки нет, проверьте права в настройках пользователя. "
        "Обычно права назначаются директору, заместителям директора и методистам.",
        size=10.5,
        color=MUTED,
        italic=True,
        after=8,
    )
    add_heading(document, "Что видно в списке", level=2)
    add_table(
        document,
        ["Колонка", "Значение"],
        [
            ("№ и дата", "Номер протокола и дата заседания."),
            ("Вид", "«Конструктор» или «Архивный Word»."),
            ("Статус", "«Черновик» или «Выпущен»."),
            ("Пункты / приложения", "Количество структурированных пунктов и прикреплённых файлов."),
            ("Добавил", "Сотрудник, создавший запись."),
            ("Действия", "Открыть, скачать Word, сделать выписку или удалить — согласно правам."),
        ],
        [2300, 7338],
        first_col_bold=True,
    )

    add_figure(
        document,
        "00-protocol-list-cropped.png",
        "Рисунок 1. Список протоколов выбранного учебного года",
        "Список протоколов педагогических советов: номер, дата, вид, статус, количество пунктов и приложений, автор и доступные действия.",
        width_mm=168,
    )

    add_page_break(document)

    # 3. Create protocol requisites.
    add_page_heading(
        document,
        "Раздел 3",
        "Создание протокола",
        "Сначала заполняются общие реквизиты заседания.",
    )
    add_numbered_steps(
        document,
        [
            ("Нажмите «Создать протокол»", "Откроется форма с предварительным просмотром справа."),
            ("Заполните реквизиты", "Поля со звёздочкой и обязательные данные должны быть заполнены."),
            ("Проверьте бланк", "В предварительном просмотре должно отображаться название нужной школы."),
        ],
        decimal_abstract,
    )
    add_heading(document, "Как заполнять реквизиты", level=2)
    add_table(
        document,
        ["Поле", "Как заполнять", "Пример"],
        [
            ("Учебный год", "Выбирается для хранения протокола; после создания не меняется.", "2025/2026"),
            ("Номер протокола", "Номер по журналу педагогических советов.", "8"),
            ("Дата", "Дата фактического заседания в пределах выбранного учебного года.", "31.03.2026"),
            ("Время начала", "Время начала всего заседания, а не длительность доклада.", "14:30"),
            ("Число присутствующих", "Только количество, без поимённого списка.", "56"),
            ("Председатель", "Сначала должность, затем ФИО из кадров.", "Директор — Жданова И.Д."),
            ("Секретарь", "Сначала должность, затем ФИО из кадров.", "Методист — Бочарова А.С."),
        ],
        [2250, 4600, 2788],
        font_size=9.7,
        first_col_bold=True,
    )
    add_callout(
        document,
        "Поиск сотрудника",
        "В поле ФИО начните печатать фамилию. Список сократится автоматически; выберите нужного "
        "человека. Основная должность берётся из кадров, но роль в документе указывается отдельно.",
        fill=LIGHT_BLUE,
    )
    add_heading(document, "Пример заполненного блока", level=2)
    add_table(
        document,
        ["Реквизит", "Значение"],
        [
            ("Учебный год", "2025/2026"),
            ("Протокол", "№ 8 от 31.03.2026, начало в 14:30"),
            ("Присутствовали", "56 человек"),
            ("Председатель", "Директор Жданова И.Д."),
            ("Секретарь", "Методист Бочарова А.С."),
        ],
        [2850, 6788],
        first_col_bold=True,
    )
    add_callout(
        document,
        "Важно",
        "Не путайте «Время начала» заседания и «Продолжительность выступления» внутри пункта. "
        "Для выступления по умолчанию устанавливается 10 минут.",
        fill="FFF8E8",
        label_color=CAUTION,
    )

    add_figure(
        document,
        "01-requisites-and-preview.png",
        "Рисунок 2. Реквизиты заседания и предварительный просмотр протокола",
        "Форма нового протокола с учебным годом, номером, датой, временем, числом присутствующих, председателем и секретарём; справа показан предварительный просмотр.",
    )

    add_page_break(document)

    # 4. Agenda item, attachment, voting.
    add_page_heading(
        document,
        "Раздел 4",
        "Добавление пунктов повестки",
        "Каждый вопрос оформляется отдельной карточкой.",
    )
    add_heading(document, "Порядок заполнения пункта", level=2)
    add_numbered_steps(
        document,
        [
            ("Укажите вопрос повестки", "Сформулируйте коротко, о чём докладывают и что решает совет."),
            ("Проверьте длительность", "По умолчанию — 10 минут; при необходимости измените."),
            ("Выберите докладчика", "Должность и ФИО выбираются отдельно; фамилию можно найти набором."),
            ("Запишите содержание", "Кратко изложите сведения, выводы и предложения докладчика."),
            ("Сформулируйте решение", "Запишите принятое решение полностью и однозначно."),
            ("Добавьте приложения", "Прикрепите .docx к этому пункту и используйте показанный номер."),
            ("Введите голоса", "Укажите «за», «против» и «воздержались» количественно."),
        ],
        decimal_abstract,
    )
    add_heading(document, "Подсказки для текста", level=2)
    add_table(
        document,
        ["Поле", "Удачная формулировка", "Не следует писать"],
        [
            (
                "Содержание выступления",
                "«представила результаты готовности обучающихся к ГИА, сведения об отсутствии академической задолженности и предложила допустить обучающихся к экзаменам.»",
                "«Доклад», «Выступила», одно-два общих слова.",
            ),
            (
                "Решили",
                "«Допустить обучающихся 9-х классов к ГИА в форме ОГЭ согласно приложению № 1.»",
                "«Принято», «Решение положительное» без содержания решения.",
            ),
        ],
        [2150, 4600, 2888],
        font_size=9.3,
        first_col_bold=True,
    )
    add_heading(document, "Полный пример пункта", level=2)
    add_table(
        document,
        ["Поле", "Пример"],
        [
            ("Вопрос повестки", "О допуске обучающихся к ГИА в форме ОГЭ"),
            ("Продолжительность", "10 минут"),
            ("Докладчик", "Заместитель директора Архангельская Т.М."),
            (
                "Содержание",
                "представила сведения о готовности обучающихся, результатах освоения образовательных "
                "программ и отсутствии академической задолженности.",
            ),
            (
                "Решили",
                "Допустить обучающихся 9-х классов к ГИА в форме ОГЭ согласно приложению № 1.",
            ),
            ("Приложение", "Список обучающихся.docx — приложение № 1"),
            ("Голосование", "за — 54; против — 0; воздержались — 2"),
        ],
        [2550, 7088],
        font_size=9.7,
        first_col_bold=True,
    )

    add_figure(
        document,
        "02-agenda-item-and-preview.png",
        "Рисунок 3. Заполнение пункта повестки и результат в предварительном просмотре",
        "Карточка пункта повестки с продолжительностью, должностью и ФИО докладчика, содержанием выступления и решением; справа показан сформированный текст.",
    )

    add_page_break(document)

    # 5. Attachments, votes and generated text.
    add_page_heading(
        document,
        "Раздел 5",
        "Приложения и голосование",
        "Приложение связано с конкретным пунктом и печатается вместе с ним.",
    )
    add_heading(document, "Работа с приложениями", level=2)
    add_bullets(
        document,
        [
            "загружается только файл Word формата .docx;",
            "файл выбирается внутри того пункта, к которому относится;",
            "система показывает следующий свободный номер приложения;",
            "до сохранения файл отмечается как ожидающий загрузки;",
            "после нажатия «Сохранить» файл передаётся на сервер;",
            "в полном протоколе и выписке приложения добавляются в конец единого Word-файла.",
        ],
        bullet_abstract,
    )
    add_callout(
        document,
        "Ссылка в решении",
        "Номер показывается до сохранения, поэтому его можно сразу написать в тексте: "
        "«…согласно приложению № 1». Если загружается несколько файлов, проверьте их номера перед выпуском.",
        fill=LIGHT_BLUE,
    )
    add_heading(document, "Контроль голосов", level=2)
    add_text_paragraph(
        document,
        "Система сравнивает сумму голосов с числом присутствующих и показывает остаток. "
        "Сумма может быть меньше числа присутствующих, но не может его превышать.",
        after=6,
    )
    add_table(
        document,
        ["Присутствуют", "За", "Против", "Воздержались", "Подсказка"],
        [
            ("56", "54", "0", "0", "Осталось распределить 2 голоса"),
            ("56", "54", "0", "2", "Все 56 голосов распределены"),
            ("56", "57", "0", "0", "Ошибка: превышение на 1 голос"),
        ],
        [1550, 1150, 1350, 1700, 3888],
        font_size=9.5,
    )
    add_heading(document, "Как пункт выглядит в Word", level=2)
    add_code_sample(
        document,
        [
            "I. Слушали: Заместителя директора Архангельскую Т.М.,",
            "представившую сведения о готовности обучающихся к ГИА.",
            "",
            "Решили: допустить обучающихся 9-х классов к ГИА",
            "в форме ОГЭ согласно приложению № 1.",
            "",
            "Приложение № 1 к пункту 1.",
            "Голосовали: за — 54, против — 0, воздержались — 2.",
        ],
        "Условный фрагмент сформированного документа",
    )
    add_callout(
        document,
        "Если загрузка не завершилась",
        "Протокол сохраняется как черновик и не выпускается. Исправьте файл или повторите загрузку, "
        "затем снова нажмите «Сохранить и выпустить».",
        fill="FFF8E8",
        label_color=CAUTION,
    )

    add_figure(
        document,
        "03-attachment-voting-and-preview.png",
        "Рисунок 4. Приложение к пункту, голосование и итоговый текст",
        "К пункту прикреплено приложение номер один; голоса распределены как 54 за, 0 против и 2 воздержались при 56 присутствующих.",
    )

    add_page_break(document)

    # 6. Status, release and concurrent editing.
    add_page_heading(
        document,
        "Раздел 6",
        "Сохранение, выпуск и совместная работа",
        "В модуле используются только два понятных статуса.",
    )
    add_table(
        document,
        ["Статус", "Назначение", "Дальнейшие действия"],
        [
            (
                "Черновик",
                "Рабочая версия. Можно добавлять и исправлять пункты и приложения.",
                "«Сохранить черновик» или «Сохранить и выпустить».",
            ),
            (
                "Выпущен",
                "Версия подготовлена для скачивания, печати и бумажного подписания.",
                "При изменении нажмите «Сохранить и перевыпустить».",
            ),
        ],
        [1750, 4300, 3588],
        font_size=9.7,
        first_col_bold=True,
    )
    add_heading(document, "Как правильно завершить работу", level=2)
    add_numbered_steps(
        document,
        [
            ("Проверьте предварительный просмотр", "Убедитесь, что школа, реквизиты и пункты отображаются верно."),
            ("Сохраните", "Для незавершённой работы выберите «Сохранить черновик»."),
            ("Выпустите", "Для готового документа выберите «Сохранить и выпустить»."),
            ("Скачайте Word", "В списке протоколов нажмите «Скачать Word»."),
            ("Распечатайте и подпишите", "Подписи ставятся на бумажном экземпляре."),
        ],
        decimal_abstract,
    )
    add_callout(
        document,
        "Перевыпуск",
        "Выпущенный протокол не блокируется навсегда. Откройте его, внесите изменения и нажмите "
        "«Сохранить и перевыпустить». Система соберёт актуальную полную версию.",
        fill=LIGHT_BLUE,
    )
    add_heading(document, "Если работают два человека", level=2)
    add_table(
        document,
        ["Ситуация", "Что произойдёт"],
        [
            (
                "Пользователи изменяют разные пункты",
                "Изменения объединятся автоматически. Каждый сохранённый пункт попадёт в актуальный протокол.",
            ),
            (
                "Пользователи изменяют один и тот же пункт",
                "Сохранится первая правка. Второй пользователь получит сообщение о конфликте; чужой текст не перезапишется.",
            ),
            (
                "Одновременно меняются общие реквизиты",
                "Система сообщит о конфликте реквизитов. Нужно открыть актуальную версию и повторить изменение.",
            ),
        ],
        [3400, 6238],
        font_size=9.7,
        first_col_bold=True,
    )
    add_callout(
        document,
        "Разрешение конфликта",
        "Скопируйте свой текст, закройте окно, снова откройте протокол, проверьте последнюю версию "
        "и внесите изменение повторно. Не обновляйте страницу, пока не скопировали несохранённый текст.",
        fill="FFF8E8",
        label_color=CAUTION,
    )

    add_page_break(document)

    # 7. Extract.
    add_page_heading(
        document,
        "Раздел 7",
        "Формирование выписки",
        "Выписка создаётся из выбранных пунктов уже введённого протокола.",
    )
    add_numbered_steps(
        document,
        [
            ("Найдите протокол", "В списке нажмите «Выписка»."),
            ("Выберите пункты", "Отметьте один или несколько пунктов. Рядом видны связанные приложения."),
            ("Проверьте заверителя", "По умолчанию указан текущий пользователь; при необходимости замените."),
            ("Добавьте заверителей", "Можно указать нескольких сотрудников, у каждого отдельно должность и ФИО."),
            ("Определите информационные данные", "Председатель и секретарь по умолчанию включаются без строк для подписи."),
            ("При необходимости включите утверждение", "Поставьте флаг и выберите должность и ФИО утверждающего."),
            ("Скачайте", "Нажмите «Скачать выписку Word»."),
        ],
        decimal_abstract,
    )
    add_heading(document, "Что попадёт в выписку", level=2)
    add_bullets(
        document,
        [
            "реквизиты исходного протокола;",
            "полный текст выбранных пунктов: «Слушали», «Решили», голосование;",
            "все приложения, прикреплённые к выбранным пунктам;",
            "председатель и секретарь информационно — если флаг оставлен включённым;",
            "строки «Верно», должность, место подписи и ФИО заверителя;",
            "место печати «М.П.» на каждой выписке;",
            "блок «УТВЕРЖДАЮ» — только если включено отдельное утверждение.",
        ],
        bullet_abstract,
    )
    add_heading(document, "Пример", level=2)
    add_table(
        document,
        ["Настройка", "Выбранное значение", "Результат"],
        [
            (
                "Пункт",
                "Пункт 3. Допуск обучающихся к ГИА в форме ОГЭ",
                "В выписку попадёт полный текст пункта 3.",
            ),
            (
                "Приложение",
                "К пункту 3 прикреплено приложение № 1",
                "Приложение № 1 автоматически добавится в тот же Word-файл.",
            ),
            (
                "Заверитель",
                "Методист Петрова П.П.",
                "Появится строка для бумажной подписи и «М.П.».",
            ),
            (
                "Отдельное утверждение",
                "Не требуется",
                "Блок утверждения не печатается.",
            ),
        ],
        [2200, 4000, 3438],
        font_size=9.4,
        first_col_bold=True,
    )
    add_callout(
        document,
        "Замена заверителя",
        "Автоматически выбранного сотрудника можно заменить вручную. Должность обычно подтягивается "
        "из кадров, но в форме выписки её можно выбрать отдельно.",
        fill=LIGHT_BLUE,
    )
    add_callout(
        document,
        "Ограничение",
        "Автоматическая выписка доступна только для протоколов, созданных в конструкторе. "
        "Из архивного Word-файла система не может определить пункты и приложения.",
        fill="FFF8E8",
        label_color=CAUTION,
    )

    add_figure(
        document,
        "04-extract-selection-cropped.png",
        "Рисунок 5. Выбор пункта и заверителя при формировании выписки",
        "Форма формирования выписки: выбран пункт протокола с приложением номер один, автоматически указан сотрудник, формирующий выписку, и включены информационные данные председателя и секретаря.",
        width_mm=145,
    )

    add_page_break(document)

    # 8. Old Word archive and deletion.
    add_page_heading(
        document,
        "Раздел 8",
        "Старые протоколы и удаление",
        "Готовые документы прошлых лет можно хранить без переноса в конструктор.",
    )
    add_heading(document, "Загрузка старого протокола", level=2)
    add_numbered_steps(
        document,
        [
            ("Нажмите «Загрузить старый Word»", "Откроется отдельное окно."),
            ("Выберите учебный год", "Укажите год, в котором проходил педагогический совет."),
            ("Введите номер и дату", "Дата должна попадать в период с 1 августа по 31 июля."),
            ("Выберите файл", "Поддерживается готовый документ Word .docx."),
            ("Сохраните", "Файл появится в списке как «Архивный Word» со статусом «Выпущен»."),
        ],
        decimal_abstract,
    )
    add_table(
        document,
        ["Возможность", "Конструктор", "Архивный Word"],
        [
            ("Хранение и скачивание", "Да", "Да"),
            ("Редактирование отдельных пунктов", "Да", "Нет"),
            ("Добавление приложений по пунктам", "Да", "Нет"),
            ("Автоматическая выписка", "Да", "Нет"),
            ("Единый сформированный Word", "Да", "Хранится исходный загруженный файл"),
        ],
        [3900, 2400, 3338],
        font_size=9.7,
        first_col_bold=True,
    )
    add_callout(
        document,
        "Формат файла",
        "Используйте .docx. Старый формат .doc предварительно откройте в Word и сохраните как "
        "«Документ Word (*.docx)». Максимальный размер файла — 30 МБ.",
        fill=CALLOUT,
    )
    add_figure(
        document,
        "05-old-word-upload-cropped.png",
        "Рисунок 6. Загрузка старого готового протокола Word",
        "Окно загрузки старого протокола с выбором учебного года, вводом номера и даты педагогического совета и выбором файла Word.",
        width_mm=145,
    )
    add_heading(document, "Удаление протокола", level=2)
    add_numbered_steps(
        document,
        [
            ("Нажмите «Удалить»", "Кнопка находится в строке протокола и доступна при праве редактирования."),
            ("Прочитайте предупреждение", "Система перечислит удаляемые данные."),
            ("Подтвердите", "Будут удалены протокол, пункты, приложения и сохранённый Word-файл."),
        ],
        decimal_abstract,
    )
    add_callout(
        document,
        "Необратимое действие",
        "Восстановить удалённый протокол через интерфейс нельзя. Перед удалением скачайте нужный Word-файл "
        "и убедитесь, что выбран правильный учебный год и номер.",
        fill="FDECEC",
        label_color=RISK,
    )

    add_page_break(document)

    # 9. Troubleshooting and checklists.
    add_page_heading(
        document,
        "Раздел 9",
        "Ошибки и контрольные списки",
        "Что проверить перед обращением к администратору.",
    )
    add_heading(document, "Частые ситуации", level=2)
    add_table(
        document,
        ["Сообщение или проблема", "Что сделать"],
        [
            (
                "Дата не относится к учебному году",
                "Проверьте границы: с 1 августа первого года по 31 июля следующего.",
            ),
            (
                "Голосов больше числа присутствующих",
                "Исправьте «за», «против» или «воздержались»; сумма не должна превышать общее число.",
            ),
            (
                "Сотрудник не находится",
                "Начните вводить фамилию; если записи нет, проверьте кадровый справочник и статус сотрудника.",
            ),
            (
                "Word-файл не загружается",
                "Проверьте расширение .docx, размер до 30 МБ и повторите сохранение.",
            ),
            (
                "Нет кнопки «Выписка» или «Скачать Word»",
                "Проверьте право экспорта для модуля «Педагогические советы».",
            ),
            (
                "Пункт изменил другой пользователь",
                "Скопируйте свой текст, заново откройте протокол и повторите правку в актуальной версии.",
            ),
            (
                "Выпуск не завершился",
                "Черновик уже сохранён. Исправьте указанную причину и снова нажмите «Сохранить и выпустить».",
            ),
        ],
        [3300, 6338],
        font_size=9.4,
        first_col_bold=True,
    )
    add_heading(document, "Проверка перед выпуском протокола", level=2)
    add_bullets(
        document,
        [
            "выбран правильный учебный год, номер, дата и время начала;",
            "указаны число присутствующих, председатель и секретарь;",
            "у каждого пункта есть вопрос, длительность, докладчик, содержание и решение;",
            "в решениях стоят правильные ссылки на номера приложений;",
            "все приложения отображаются как загруженные, а не ожидающие;",
            "голоса по каждому пункту не превышают число присутствующих;",
            "предварительный просмотр показывает нужную школу и правильный текст;",
            "после выпуска скачан Word-файл для печати и бумажного подписания.",
        ],
        bullet_abstract,
    )
    add_heading(document, "Проверка перед скачиванием выписки", level=2)
    add_bullets(
        document,
        [
            "отмечены только нужные пункты;",
            "проверено, какие приложения связаны с выбранными пунктами;",
            "выбран правильный заверитель и его должность;",
            "председатель и секретарь включены информационно, если это требуется;",
            "отдельное утверждение включено только при реальной необходимости;",
            "для утверждения выбраны должность и ФИО утверждающего;",
            "в готовом Word-файле присутствуют строки для подписи и «М.П.».",
        ],
        bullet_abstract,
    )
    add_callout(
        document,
        "Результат",
        "После заполнения одного протокола система хранит его по учебному году, формирует полный Word-документ "
        "и позволяет получать выписки без повторного набора текста.",
        fill=LIGHT_BLUE,
        label_color=POSITIVE,
    )
    add_text_paragraph(
        document,
        "Если ошибка повторяется после проверки данных, передайте администратору номер протокола, "
        "учебный год, время ошибки и текст сообщения на экране.",
        size=10,
        color=MUTED,
        italic=True,
        before=4,
        after=0,
    )

    return document


def audit_document(document: Document) -> None:
    section = document.sections[0]
    assert abs(section.page_width.twips - PAGE_WIDTH_DXA) <= 2
    assert abs(section.page_height.twips - PAGE_HEIGHT_DXA) <= 2
    assert abs(section.left_margin.twips - MARGIN_DXA) <= 2
    assert abs(section.right_margin.twips - MARGIN_DXA) <= 2
    assert document.styles["Normal"].font.name == FONT
    assert document.styles["Normal"].font.size.pt == 11
    assert document.styles["Heading 1"].font.size.pt == 16
    assert document.styles["Heading 2"].font.size.pt == 13
    assert document.styles["Heading 3"].font.size.pt == 12

    for table in document.tables:
        tblpr = table._tbl.tblPr
        tblw = tblpr.find(qn("w:tblW"))
        tblind = tblpr.find(qn("w:tblInd"))
        assert tblw is not None and int(tblw.get(qn("w:w"))) == CONTENT_WIDTH_DXA
        assert tblw.get(qn("w:type")) == "dxa"
        assert tblind is not None and int(tblind.get(qn("w:w"))) == TABLE_INDENT_DXA
        grid_widths = [
            int(col.get(qn("w:w")))
            for col in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        assert sum(grid_widths) == CONTENT_WIDTH_DXA
        for row in table.rows:
            assert len(row.cells) == len(grid_widths)
            for index, cell in enumerate(row.cells):
                tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tcw is not None
                assert int(tcw.get(qn("w:w"))) == grid_widths[index]
                assert tcw.get(qn("w:type")) == "dxa"

    numbering = document.part.numbering_part.element
    custom_abstracts = numbering.findall(qn("w:abstractNum"))
    assert len(custom_abstracts) >= 2


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    audit_document(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
