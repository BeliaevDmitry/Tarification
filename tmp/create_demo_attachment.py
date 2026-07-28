from docx import Document


document = Document()
document.add_heading("Приложение № 1", level=1)
document.add_paragraph(
    "к протоколу педагогического совета от 31.03.2026 № 8"
)
document.add_heading("Список обучающихся, допущенных к ГИА", level=2)
table = document.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.rows[0].cells[0].text = "№"
table.rows[0].cells[1].text = "ФИО обучающегося"
table.rows[0].cells[2].text = "Класс"
for number, name, school_class in (
    (1, "Иванов Иван Иванович", "9А"),
    (2, "Петрова Анна Сергеевна", "9Б"),
):
    cells = table.add_row().cells
    cells[0].text = str(number)
    cells[1].text = name
    cells[2].text = school_class

document.save(
    r"C:\Users\dimah\IdeaProjects\Tarification\tmp\Приложение_1_Список_допущенных.docx"
)
